# -*- coding: utf-8 -*-
import sqlite3
import config
import dbtool
from collections import defaultdict
from docx import Document
import time
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

config = config.Config()


class Kanon():

    def __init__(self):
        self.dbtool = dbtool.DbTool()
        # print(self.dbtool.get_books_by_seasons())

    def get_genres(self):
        return [{"genre_id": 1, "genre_name": "Próza", "color": "#148a14"},
                {"genre_id": 2, "genre_name": "Poezie", "color": "#1b2ec5"},
                {"genre_id": 3, "genre_name": "Drama", "color": "#865a15"}]

    def check_books(self, data, request):
        bookIds = []
        rules = {
            "authorMax2": True,  # max 2 dila od jednoho autora
            "min2Proza": False,
            "min2Poezie": False,
            "min2Drama": False,
            "20Books": False,
            "season1": False,
            "season2": False,
            "season3": False,
            "season4": False,
        }

        if not "checked[]" in data.keys():
            return self.process_rules_to_string(rules, 0)
        for b in request.form.getlist("checked[]"):
            bookIds.append(int(b))
        print(bookIds)
        if len(bookIds) == 0:
            return self.process_rules_to_string(rules, 0)

        rules["20Books"] = (len(bookIds) == 20)

        books = self.dbtool.get_selected_books(bookIds)
        authors = defaultdict(int)
        authorMax2 = True
        seasons = defaultdict(int)
        genres = defaultdict(int)
        for book in books:
            authors[book["author_id"]] += 1
            if authors[book["author_id"]] > 2:
                authorMax2 = False

            seasons[book["book_season"]] += 1
            genres[book["book_genre"]] += 1

        rules["min2Proza"] = genres[1] >= 2
        rules["min2Poezie"] = genres[2] >= 2
        rules["min2Drama"] = genres[3] >= 2

        rules["season1"] = seasons[1] >= 2
        rules["season2"] = seasons[2] >= 3
        rules["season3"] = seasons[3] >= 4
        rules["season4"] = seasons[4] >= 5
        rules["authorMax2"] = authorMax2

        # print(len(books))

        return self.process_rules_to_string(rules, len(bookIds))

    def process_rules_to_string(self, rules, booksCount):
        result = "<h2>Podmínky pro výběr z kánonu</h2><br/>" \
                 "<span id='allRulesOk' style='color: " + (
                     "red'>Nejsou splněny všechny podmínky!" if False in rules.values()
                     else "green'>Všechny podmínky splněny! ") + \
                 "</span><br/>"

        result += "<span style='color: " + ("green'> OK: " if rules["20Books"] else "red'> NOK: ") + \
                  "20 knih (vybráno " + str(booksCount) + ")</span><br/>"

        result += "<span style='color: " + ("green'> OK: " if rules["authorMax2"] else "red'> NOK: ") + \
                  "Maximálně 2 díla od každého autora</span><br/>"

        result += "<span style='color: " + ("green'> OK: " if rules["min2Proza"] else "red'> NOK: ") + \
                  "Alespoň 2 prózy</span><br/>"

        result += "<span style='color: " + ("green'> OK: " if rules["min2Poezie"] else "red'> NOK: ") + \
                  "Alespoň 2 poezie</span><br/>"

        result += "<span style='color: " + ("green'> OK: " if rules["min2Drama"] else "red'> NOK: ") + \
                  "Alespoň 2 dramata</span><br/>"

        result += "<span style='color: " + ("green'> OK: " if rules["season1"] else "red'> NOK: ") + \
                  "Alespoň 2 díla z období 18. století</span><br/>"

        result += "<span style='color: " + ("green'> OK: " if rules["season2"] else "red'> NOK: ") + \
                  "Alespoň 3 díla z období 19. století</span><br/>"

        result += "<span style='color: " + ("green'> OK: " if rules["season3"] else "red'> NOK: ") + \
                  "Alespoň 4 díla ze světové literatury 20. a 21. století</span><br/>"

        result += "<span style='color: " + ("green'> OK: " if rules["season4"] else "red'> NOK: ") + \
                  "Alespoň 5 děl z české literatury 20. a 21. století</span><br/>"

        return result

    def create_docx_file(self, data):

        document = Document()
        p = document.add_paragraph()
        p.add_run('Maturitní témata z českého jazyka').bold = True
        s = p.paragraph_format
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph()

        p.add_run('Jméno a příjmení: ' + data["name"]).bold = True
        p.add_run('\nTřída: ' + data["class"]).bold = True
        p.add_run('\nŠkolní rok: ' + data["year"]).bold = True
        p.add_run("\n")

        books = self.dbtool.get_selected_books(data["books"].split(","))
        season = 0
        for book in books:
            if book["book_season"] != season:
                season = book["book_season"]
                if season == 1:
                    p = document.add_paragraph()
                    p.add_run('Světová a česká literatura do konce 18. století').bold = True
                elif season == 2:
                    p = document.add_paragraph()
                    p.add_run('Světová a česká literatura 19. století').bold = True
                elif season == 3:
                    p = document.add_paragraph()
                    p.add_run('Světová literatura 20. a 21. století').bold = True
                elif season == 4:
                    p = document.add_paragraph()
                    p.add_run('Česká literatura 20. a 21. století').bold = True

            document.add_paragraph(((book["author_name"] + ": ") if not book["author_name"].startswith("skryt") else "")
                                   + book["book_name"] +
                                   " (" + str(book["book_order"]) + ")", style='List Number')

        docname = 'kanon_' + data["name"].split(" ")[-1] + "_" + data["class"] + '.doc'
        document.save(os.path.dirname(os.path.abspath(__file__)) + "/exports/" + docname)
        return docname
